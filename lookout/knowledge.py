from __future__ import annotations
from pathlib import Path
from pypdf import PdfReader
from docx import Document
from .extensions import db
from .models import KnowledgeDocument, KnowledgeChunk

ALLOWED = {".pdf", ".docx", ".txt", ".md"}


def extract_text(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        reader = PdfReader(str(path))
        return "\n".join((p.extract_text() or "") for p in reader.pages)
    if suffix == ".docx":
        doc = Document(str(path))
        return "\n".join(p.text for p in doc.paragraphs)
    if suffix in {".txt", ".md"}:
        return path.read_text(errors="ignore")
    raise ValueError("Unsupported file type")


def chunk_text(text: str, target: int = 1800, overlap: int = 200) -> list[str]:
    clean = "\n".join(line.strip() for line in text.splitlines() if line.strip())
    if not clean:
        return []
    out, start = [], 0
    while start < len(clean):
        end = min(len(clean), start + target)
        if end < len(clean):
            boundary = clean.rfind(" ", start, end)
            if boundary > start + target // 2:
                end = boundary
        out.append(clean[start:end].strip())
        if end >= len(clean): break
        start = max(start + 1, end - overlap)
    return out


def ingest(project_id: int, title: str, filename: str, path: Path) -> KnowledgeDocument:
    text = extract_text(path)
    doc = KnowledgeDocument(project_id=project_id, title=title, filename=filename, full_text=text)
    db.session.add(doc); db.session.flush()
    for i, chunk in enumerate(chunk_text(text)):
        db.session.add(KnowledgeChunk(document_id=doc.id, project_id=project_id, chunk_index=i, content=chunk))
    db.session.commit()
    return doc
